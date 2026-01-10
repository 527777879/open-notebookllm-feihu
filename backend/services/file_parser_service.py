"""
文件解析服務
"""
import os
import logging
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


class FileParserService:
    """文件解析服務"""

    SUPPORTED_EXTENSIONS = {
        'pdf': 'parse_pdf',
        'txt': 'parse_text',
        'md': 'parse_text',
        'docx': 'parse_docx',
        'doc': 'parse_docx',
        'xlsx': 'parse_excel',
        'xls': 'parse_excel',
        'csv': 'parse_csv',
    }

    def parse_file(self, file_path: str, filename: str) -> Tuple[Optional[str], Optional[str]]:
        """
        解析文件

        Args:
            file_path: 文件路徑
            filename: 文件名

        Returns:
            (提取的文字內容, 錯誤訊息)
        """
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        if ext not in self.SUPPORTED_EXTENSIONS:
            return None, f"不支援的文件格式: {ext}"

        parser_method = getattr(self, self.SUPPORTED_EXTENSIONS[ext])

        try:
            content = parser_method(file_path)
            if not content:
                return None, "無法提取文件內容"
            return content, None
        except Exception as e:
            logger.error(f"文件解析失敗: {e}")
            return None, str(e)

    def parse_pdf(self, file_path: str) -> Optional[str]:
        """解析 PDF 文件"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return '\n\n'.join(text_parts)

        except Exception as e:
            logger.error(f"PDF 解析失敗: {e}")
            raise

    def parse_text(self, file_path: str) -> Optional[str]:
        """解析純文字文件"""
        try:
            # 嘗試不同編碼
            encodings = ['utf-8', 'utf-8-sig', 'big5', 'gb2312', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            raise ValueError("無法辨識文件編碼")

        except Exception as e:
            logger.error(f"文字文件解析失敗: {e}")
            raise

    def parse_docx(self, file_path: str) -> Optional[str]:
        """解析 Word 文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 也提取表格內容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Word 文件解析失敗: {e}")
            raise

    def parse_excel(self, file_path: str) -> Optional[str]:
        """解析 Excel 文件"""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"## 工作表: {sheet_name}\n")

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_values):
                        text_parts.append(' | '.join(row_values))

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Excel 文件解析失敗: {e}")
            raise

    def parse_csv(self, file_path: str) -> Optional[str]:
        """解析 CSV 文件"""
        try:
            import csv

            text_parts = []

            # 嘗試不同編碼
            encodings = ['utf-8', 'utf-8-sig', 'big5', 'gb2312', 'latin-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, newline='') as f:
                        reader = csv.reader(f)
                        for row in reader:
                            if any(row):
                                text_parts.append(' | '.join(row))
                    break
                except UnicodeDecodeError:
                    continue

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"CSV 文件解析失敗: {e}")
            raise


# 全局實例
_file_parser: Optional[FileParserService] = None


def get_file_parser() -> FileParserService:
    """取得文件解析服務實例"""
    global _file_parser
    if _file_parser is None:
        _file_parser = FileParserService()
    return _file_parser
